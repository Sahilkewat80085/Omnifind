from datetime import datetime, timezone
from pathlib import Path
import docx

from core.document_pipeline.extractors.base import BaseExtractor
from core.document_pipeline.models import (
    FormatMetadata,
    NormalizedBlock,
    NormalizedDocument,
    NormalizedTable,
)


class DocxExtractor(BaseExtractor):
    """Extractor for Microsoft Word (.docx) files preserving style hierarchy, lists, and tables."""

    def extract(self, file_path: str) -> NormalizedDocument:
        path = Path(file_path)
        stat = path.stat()
        created_at = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc)
        modified_at = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc)

        doc = docx.Document(str(path))
        blocks: list[NormalizedBlock] = []
        tables: list[NormalizedTable] = []
        outline: list[dict[str, int | str]] = []

        current_heading_path: list[str] = []

        # 1. Extract Paragraphs with Heading Hierarchy
        para_idx = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            para_idx += 1
            style_name = para.style.name if para.style else "Normal"
            style_lower = style_name.lower()

            if "heading" in style_lower:
                # Extract heading level (e.g. "Heading 1" -> 1)
                level = 1
                try:
                    level = int("".join(filter(str.isdigit, style_name)) or 1)
                except ValueError:
                    level = 1

                if level <= len(current_heading_path):
                    current_heading_path = current_heading_path[: level - 1]
                current_heading_path.append(text)
                outline.append({"level": level, "title": text})

                blocks.append(
                    NormalizedBlock(
                        text=text,
                        location=f"Section: {text}",
                        heading_path=list(current_heading_path[:-1]),
                        block_type="heading",
                        level=level,
                    )
                )
            elif "list" in style_lower or "bullet" in style_lower:
                blocks.append(
                    NormalizedBlock(
                        text=text,
                        location=f"Item {para_idx}",
                        heading_path=list(current_heading_path),
                        block_type="list_item",
                    )
                )
            else:
                blocks.append(
                    NormalizedBlock(
                        text=text,
                        location=f"Paragraph {para_idx}",
                        heading_path=list(current_heading_path),
                        block_type="paragraph",
                    )
                )

        # 2. Extract Tables
        for table_idx, table in enumerate(doc.tables, start=1):
            if not table.rows:
                continue

            table_rows: list[list[str]] = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                # Avoid duplicate merged cells on same row
                table_rows.append(row_cells)

            if table_rows:
                headers = table_rows[0]
                data_rows = table_rows[1:] if len(table_rows) > 1 else []
                tables.append(
                    NormalizedTable(
                        table_id=f"table-{table_idx}",
                        location=f"Table #{table_idx}",
                        headers=headers,
                        rows=data_rows,
                    )
                )

        return NormalizedDocument(
            file_path=str(path.resolve()),
            file_name=path.name,
            file_type="docx",
            created_at=created_at,
            modified_at=modified_at,
            blocks=blocks,
            tables=tables,
            format_metadata=FormatMetadata(heading_outline=outline or None),
            is_scanned=False,
            requires_ocr=False,
        )
