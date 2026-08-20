import docx

from core.parsers.base import PageText


def parse_docx(path: str) -> list[PageText]:
    document = docx.Document(path)
    lines: list[str] = []

    for p in document.paragraphs:
        txt = p.text.strip()
        if txt:
            lines.append(txt)

    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                # Deduplicate identical adjacent cells caused by merged table cells
                deduped: list[str] = []
                for cell_txt in row_texts:
                    if not deduped or cell_txt != deduped[-1]:
                        deduped.append(cell_txt)
                if deduped:
                    lines.append(" | ".join(deduped))

    text = "\n".join(lines).strip()
    if not text:
        return []
    return [PageText(page_number=None, text=text)]
